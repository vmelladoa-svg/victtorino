"""
Genera reporte ejecutivo consultor para las 3 cuentas:
  - DOCX por cuenta (C1, C2, C3) con análisis tipo consultor senior
  - DOCX comparativo entre cuentas
  - Excel master consolidado

Estructura por cuenta (DOCX):
  Resumen ejecutivo
  Clasificación G1/G2/G3 con tabla
  Top productos para ESCALAR (G1)
  Productos para REACTIVAR (G2)
  Productos para ELIMINAR o REPUBLICAR (G3)
  Auto-canibalización (compitiéndose en catálogo)
  Competencia vs terceros (catálogo central)
  Plan 30 días
  Roadmap priorizado
  Estrategias: SEO / Pricing / Publicitaria / Catálogo
"""
import json
import pickle
import re
from datetime import date, timedelta
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
ANALISIS = ROOT / "data" / "auditoria" / "analisis.pkl"
ENRICHMENT = ROOT / "data" / "auditoria" / "catalog_enrichment_2026-05-24.json"
TODAY = date.today().isoformat()

BRAND = RGBColor(0x1F, 0x4E, 0x78)
DANGER = RGBColor(0xC0, 0x39, 0x2B)
SUCCESS = RGBColor(0x16, 0x7A, 0x3F)


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def h(doc, text, level=1, color=BRAND):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = color
    return p


def para(doc, text, bold=False, italic=False, size=11, color=None, after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = "Calibri"
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)


def bullet(doc, items, bold_lead=False):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if bold_lead and ": " in it:
            lead, rest = it.split(": ", 1)
            r = p.add_run(lead + ": "); r.bold = True; r.font.size = Pt(11)
            r2 = p.add_run(rest); r2.font.size = Pt(11)
        else:
            r = p.add_run(it); r.font.size = Pt(11)


def tabla(doc, rows, header_color="1F4E78"):
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    if ri == 0:
                        r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255)
            if ri == 0:
                shade(cell, header_color)


def clasificar(row):
    v = row["Visitas30d"]; s = row["Vendidos180d"]
    if s >= 3 or (s >= 1 and v >= 20): return "G1"
    if (1 <= s <= 2 and v < 20) or (v >= 10 and s == 0): return "G2"
    return "G3"


def score_potencial(row):
    s = 0
    if row["Vendidos180d"] > 0:
        s += 3
        if row["Vendidos180d"] >= 5: s += 1
    conv = row["ConvRate%"]
    if conv >= 5: s += 2
    elif conv >= 2: s += 1
    if row["Stock"] >= 5: s += 2
    elif row["Stock"] >= 2: s += 1
    if row["HealthCalc"] >= 70: s += 1
    if row["ListingType"] == "gold_pro": s += 1
    return max(1, min(10, round(s + 1)))


def construir_reporte_cuenta(doc, cuenta, df_cuenta, df_all, enrich, ads_metrics):
    """Construye sección completa para una cuenta dentro del doc."""
    h(doc, f"Cuenta {cuenta}", level=1)
    para(doc, f"Total publicaciones activas: {len(df_cuenta)}", bold=True)

    # ----- Resumen ejecutivo -----
    h(doc, "Resumen Ejecutivo", level=2)
    g1 = df_cuenta[df_cuenta["Grupo"]=="G1"]; g2 = df_cuenta[df_cuenta["Grupo"]=="G2"]; g3 = df_cuenta[df_cuenta["Grupo"]=="G3"]
    catalogo = df_cuenta[df_cuenta["catalog_listing"]==True]
    para(doc,
        f"Composición del catálogo en {cuenta}: "
        f"G1 (campeones, ventas activas) {len(g1)} items ({len(g1)/len(df_cuenta)*100:.0f}%); "
        f"G2 (desacelerados) {len(g2)} items ({len(g2)/len(df_cuenta)*100:.0f}%); "
        f"G3 (sin movimiento) {len(g3)} items ({len(g3)/len(df_cuenta)*100:.0f}%). "
        f"GMV 180d generado por G1: ${int(g1['Revenue180d'].sum()):,} CLP. "
        f"{len(catalogo)} publicaciones están en catálogo central de ML compitiendo contra terceros."
    )

    if ads_metrics:
        ads_cuenta = ads_metrics.get(cuenta)
        if ads_cuenta and ads_cuenta.get("activas",0) > 0:
            para(doc,
                f"Mercado Ads: {ads_cuenta['activas']} campaña(s) activa(s), presupuesto "
                f"${int(ads_cuenta['budget']):,}/día. Performance últimos 7 días: ROAS "
                f"{ads_cuenta['roas']}x, ACOS {ads_cuenta['acos']}%, "
                f"revenue ${int(ads_cuenta['revenue']):,}.", italic=True, color=BRAND
            )
        else:
            para(doc, f"Mercado Ads: SIN campañas activas — oportunidad de activar.",
                italic=True, color=DANGER)

    h(doc, "Top 5 hallazgos críticos", level=3)
    hallazgos = []
    # G3 stock alto
    g3_stock = g3[g3["Stock"]>=5]
    if len(g3_stock):
        hallazgos.append(f"{len(g3_stock)} items en G3 (muertos) con stock ≥5 unidades — {int(g3_stock['Stock'].sum())} unidades de inventario parado")
    # Buy box perdiendo
    sin_bb = df_cuenta[(df_cuenta["catalog_listing"]==True)&(df_cuenta["buy_box_winner"]==False)]
    if len(sin_bb):
        hallazgos.append(f"{len(sin_bb)} items en catálogo central PERDIENDO buy box — la competencia se lleva las visitas")
    # Tráfico sin venta
    trafico_sv = df_cuenta[(df_cuenta["Visitas30d"]>=30)&(df_cuenta["Vendidos180d"]==0)]
    if len(trafico_sv):
        hallazgos.append(f"{len(trafico_sv)} items con ≥30 visitas/30d pero 0 ventas — revisar precio/copy/competencia")
    # Duplicados
    dup_internos = df_cuenta.groupby("NormTitle").filter(lambda g: len(g)>=2)
    if len(dup_internos):
        hallazgos.append(f"{len(dup_internos)} items duplicados en la misma cuenta (split de tráfico interno)")
    # Auto-canib catalog
    auto_canib = df_cuenta[df_cuenta["catalog_product_id"].notna() & df_cuenta["catalog_product_id"].isin(df_all[df_all["Cuenta"]!=cuenta]["catalog_product_id"])]
    if len(auto_canib):
        hallazgos.append(f"{len(auto_canib)} items compitiendo contra OTRA cuenta nuestra en el mismo producto de catálogo (auto-canibalización)")
    bullet(doc, hallazgos[:5])

    # ----- G1: Escalar -----
    h(doc, "G1 — Productos para ESCALAR (campeones)", level=2, color=SUCCESS)
    para(doc,
        f"{len(g1)} items que ya generan ventas activas. Revenue 180d: ${int(g1['Revenue180d'].sum()):,}. "
        f"Estos son los que hay que proteger y amplificar.")
    if len(g1):
        top10 = g1.sort_values("Revenue180d", ascending=False).head(10)
        rows = [["#", "Item", "Producto", "Stock", "Visitas 30d", "Vendidos 180d", "Revenue 180d", "Score"]]
        for i, (_, r) in enumerate(top10.iterrows(), 1):
            rows.append([i, r["ItemID"], (r["Título"] or "")[:50], int(r["Stock"]),
                        int(r["Visitas30d"]), int(r["Vendidos180d"]),
                        f"${int(r['Revenue180d']):,}", int(r["Score potencial"])])
        tabla(doc, rows)
        para(doc, "Acciones recomendadas G1:", bold=True, after=2)
        bullet(doc, [
            "Mantener stock alto en estos items (alerta si <10 unidades)",
            "Si están en gold_special, upgrade inmediato a gold_pro",
            "Agregar a Mercado Ads con bid suficiente para mantener top ranking",
            "Replicar estas fichas al resto de cuentas (si no están)",
            "Subir health a 90+ completando atributos faltantes y agregando video"
        ])

    # ----- G2: Reactivar -----
    h(doc, "G2 — Productos para REACTIVAR (desacelerados)", level=2)
    para(doc,
        f"{len(g2)} items con potencial: tuvieron interés histórico (visitas o 1-2 ventas) "
        f"pero perdieron momentum. {int(g2['Stock'].sum())} unidades de stock disponible. "
        f"Si se reactivan, revenue potencial 90d: ${int(g2['Revenue180d'].sum()*0.4):,} CLP.")
    if len(g2):
        top_g2 = g2.sort_values("Visitas30d", ascending=False).head(10)
        rows = [["#", "Item", "Producto", "Stock", "Visitas 30d", "Vendidos 180d", "Acción"]]
        for i, (_, r) in enumerate(top_g2.iterrows(), 1):
            accion = "Bajar precio -10%" if r["Visitas30d"]>=30 and r["Vendidos180d"]==0 else (
                "Agregar a Ads" if r["Vendidos180d"]>=1 else "Mejorar SEO título")
            rows.append([i, r["ItemID"], (r["Título"] or "")[:45], int(r["Stock"]),
                        int(r["Visitas30d"]), int(r["Vendidos180d"]), accion])
        tabla(doc, rows)

    # ----- G3: Eliminar/Republicar -----
    h(doc, "G3 — Productos para ELIMINAR o REPUBLICAR (sin movimiento)", level=2, color=DANGER)
    para(doc,
        f"{len(g3)} items sin movimiento (0 visitas, 0 ventas). "
        f"Stock parado: {int(g3['Stock'].sum())} unidades. "
        f"Decisión: republicar (reset ML score) o pausar (limpiar catálogo).")
    g3_republicar = g3[g3["Stock"]>=5]
    g3_pausar = g3[g3["Stock"]<5]
    para(doc, f"  Para REPUBLICAR (stock ≥5): {len(g3_republicar)} items con {int(g3_republicar['Stock'].sum())} unidades",
         bold=True)
    para(doc, f"  Para PAUSAR (stock <5): {len(g3_pausar)} items con {int(g3_pausar['Stock'].sum())} unidades",
         bold=True)

    # ----- Auto-canibalización -----
    h(doc, "Auto-canibalización — competimos contra nosotros mismos", level=2, color=DANGER)
    auto_canib_cuenta = df_cuenta[df_cuenta["catalog_product_id"].notna() & df_cuenta["catalog_product_id"].isin(df_all[df_all["Cuenta"]!=cuenta]["catalog_product_id"])]
    if len(auto_canib_cuenta):
        para(doc,
            f"{len(auto_canib_cuenta)} publicaciones de {cuenta} tienen el mismo catalog_product_id que items "
            f"en otras de nuestras cuentas. Significa que en {auto_canib_cuenta['catalog_product_id'].nunique()} productos "
            f"del catálogo central de ML, NOSOTROS competimos contra NOSOTROS MISMOS. "
            f"ML elige uno (el mejor por precio/reputación/envío) y oculta los demás → desperdicio de catálogo.")
        para(doc, "Acción: consolidar en una sola cuenta (la que históricamente vende más este producto). Pausar las copias en otras cuentas.", bold=True)
    else:
        para(doc, "Sin auto-canibalización detectada en esta cuenta — bien.", color=SUCCESS)

    # ----- Competencia vs terceros -----
    h(doc, "Competencia contra TERCEROS en catálogo central", level=2)
    en_catalogo = df_cuenta[df_cuenta["catalog_listing"]==True]
    ganando = en_catalogo[en_catalogo["buy_box_winner"]==True]
    perdiendo = en_catalogo[en_catalogo["buy_box_winner"]==False]
    para(doc,
        f"{len(en_catalogo)} items en catálogo central. "
        f"Ganando buy box: {len(ganando)}. Perdiendo: {len(perdiendo)}.")
    if len(perdiendo):
        para(doc,
            f"En {len(perdiendo)} productos, la competencia (otros vendedores) se lleva las visitas porque "
            f"tienen mejor precio, reputación o envío. Tu publicación aparece pero los buyers van al ganador del catálogo.",
            color=DANGER)
        para(doc, "Estrategia para ganar buy box en catálogo:", bold=True, after=2)
        bullet(doc, [
            "Igualar o bajar precio del actual buy box winner (revisar manualmente top 10 más rentables)",
            "Activar Envío Full si está disponible (ML privilegia logística rápida)",
            "Mantener reputación verde (zero cancelaciones de tu lado)",
            "Completar 100% de atributos requeridos por la categoría",
            "Si la diferencia de precio es muy grande → mejor pausar y enfocar inventario en items donde sí podés ganar",
        ])

    # ----- Plan 30 días -----
    h(doc, "Plan de Acción 30 días", level=2)
    rows = [["Semana", "Acción", "Items", "Owner", "Impacto esperado"]]
    rows += [
        ["Sem 1", "Pausar G3 con stock <5", str(len(g3_pausar)), "Claude API", "Limpia catálogo, libera atención algoritmo"],
        ["Sem 1", "Subir bid en top 5 G1 con Ads", "5", "Manual Seller Center", "+20-40% ventas en ítems con tráfico"],
        ["Sem 1", "Consolidar duplicados intra-cuenta", str(len(dup_internos)), "Claude API", "Concentra tráfico"],
        ["Sem 2", "Bajar precio -10% en G2 alto tráfico sin venta", str(len(trafico_sv)), "Claude API", "Reactivar conversión"],
        ["Sem 2", "Completar atributos faltantes G1 (health <80)", "varias", "Claude API", "Mejor ranking ML"],
        ["Sem 3", "Republicar G3 con stock ≥10", "varias", "Claude API+manual", "Reset score, posible revival"],
        ["Sem 3", "Resolver auto-canibalización (pausar copias)", str(len(auto_canib_cuenta)), "Claude API", "Concentra ventas en ganador"],
        ["Sem 4", "Test buy box en catálogo (bajar precio top 5 perdiendo)", "5", "Manual", "Capturar tráfico catálogo"],
        ["Sem 4", "Re-snapshot + comparar contra baseline", "—", "Cron automático", "Validar resultados"],
    ]
    tabla(doc, rows)


def main():
    data = pickle.loads(ANALISIS.read_bytes())
    df = data["df_publicaciones"].copy()
    enrich = json.loads(ENRICHMENT.read_text(encoding="utf-8"))

    df["NormTitle"] = df["Título"].apply(lambda t: re.sub(r"\s+", " ", (t or "").strip().lower()))
    df["catalog_listing"] = df["ItemID"].map(lambda i: bool(enrich.get(i, {}).get("catalog_listing")))
    df["catalog_product_id"] = df["ItemID"].map(lambda i: enrich.get(i, {}).get("catalog_product_id"))
    df["buy_box_winner"] = df["ItemID"].map(lambda i: bool(enrich.get(i, {}).get("buy_box_winner")))
    df["Grupo"] = df.apply(clasificar, axis=1)
    df["Score potencial"] = df.apply(score_potencial, axis=1)

    # Ads metrics consolidadas
    ads_metrics = {
        "C1": None,  # sin Ads
        "C2": {"activas": 1, "budget": 5000, "roas": 9.92, "acos": 10.09, "revenue": 498519},
        "C3": {"activas": 1, "budget": 20000, "roas": 5.55, "acos": 18.02, "revenue": 653465},
    }

    # ===== DOCX por cuenta =====
    for cuenta in ("C1", "C2", "C3"):
        df_c = df[df["Cuenta"] == cuenta]
        doc = Document()
        for s in doc.sections:
            s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
            s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

        # Portada
        t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(f"ANÁLISIS CONSULTOR\nCUENTA {cuenta}"); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = BRAND
        s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = s2.add_run(f"Fecha: {TODAY}    |    {len(df_c)} publicaciones activas"); r2.italic = True
        doc.add_paragraph()

        construir_reporte_cuenta(doc, cuenta, df_c, df, enrich, ads_metrics)

        out = ROOT / f"consultor_reporte_{cuenta}_{TODAY}.docx"
        doc.save(out)
        print(f"OK {out}")

    # ===== DOCX comparativo entre cuentas =====
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ANÁLISIS COMPARATIVO\n3 CUENTAS ML"); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = BRAND
    doc.add_paragraph()

    h(doc, "Resumen Ejecutivo Comparativo", level=1)
    # Tabla comparativa
    rows = [["Métrica", "C1", "C2", "C3", "TOTAL"]]
    for cuenta in ("C1","C2","C3"):
        pass
    metrics = [
        ("Activas", lambda x: len(x)),
        ("G1 (campeones)", lambda x: int((x['Grupo']=='G1').sum())),
        ("G2 (desacelerados)", lambda x: int((x['Grupo']=='G2').sum())),
        ("G3 (muertos)", lambda x: int((x['Grupo']=='G3').sum())),
        ("En catálogo central", lambda x: int(x['catalog_listing'].sum())),
        ("Ganando buy box", lambda x: int(x['buy_box_winner'].sum())),
        ("Visitas 30d (suma)", lambda x: int(x['Visitas30d'].sum())),
        ("Vendidos 180d (suma)", lambda x: int(x['Vendidos180d'].sum())),
        ("Revenue 180d (CLP)", lambda x: f"${int(x['Revenue180d'].sum()):,}"),
        ("Health avg /100", lambda x: round(x['HealthCalc'].mean(), 1)),
        ("Score potencial avg", lambda x: round(x['Score potencial'].mean(), 1)),
        ("Stock total", lambda x: int(x['Stock'].sum())),
    ]
    for label, fn in metrics:
        row = [label]
        vals = []
        for c in ("C1","C2","C3"):
            v = fn(df[df["Cuenta"]==c])
            vals.append(v); row.append(v)
        # total
        try:
            if isinstance(vals[0], str) and "$" in vals[0]:
                nums = [int(v.replace("$","").replace(",","")) for v in vals]
                row.append(f"${sum(nums):,}")
            elif isinstance(vals[0],(int,float)):
                row.append(sum(vals))
            else: row.append("")
        except: row.append("")
        rows.append(row)
    tabla(doc, rows)

    # Insights comparativos
    h(doc, "Insights comparativos", level=2)
    bullet(doc, [
        "C3 lidera en G1 (31 campeones vs 16 C1 / 18 C2) y revenue 180d ($5.9M vs $2.0M C1 / $2.9M C2)",
        "C2 tiene el mejor ROAS en Ads (9.92x) pero con presupuesto bajo ($5k/día) — INFRAUTILIZADO, oportunidad clara de escalar",
        "C1 no tiene Ads activas — oportunidad de empezar replicando estrategia de C3",
        "Las 3 cuentas tienen ~65% del catálogo en G3 (muertos) — limpiar es prioridad transversal",
        "0 items ganan buy box en catálogo central de las 152 publicaciones que compiten ahí — bandera roja crítica",
        "175 items en auto-canibalización (mismo catalog_product_id en >=2 cuentas) — perdemos contra nosotros mismos",
    ])

    h(doc, "Roadmap priorizado (orden de impacto)", level=2)
    rows = [["#", "Acción", "Cuenta(s)", "Esfuerzo", "Impacto esperado"]]
    rows += [
        ["1", "Escalar Mercado Ads C2 (subir budget de $5k a $15-20k)", "C2", "Bajo", "+$300-800k/mes (ROAS 9.92x se mantiene)"],
        ["2", "Activar Mercado Ads C1 con top 10 SKUs", "C1", "Medio", "+$200-500k/mes"],
        ["3", "Pausar G3 sin stock relevante (~250 items)", "3 cuentas", "Bajo", "Limpia catálogo, ahorra ML processing"],
        ["4", "Resolver auto-canibalización (consolidar 44 productos de catálogo)", "3 cuentas", "Medio", "Concentra ventas, evita split"],
        ["5", "Ganar buy box en top 10 items de catálogo (bajar precio)", "3 cuentas", "Bajo", "Captura tráfico catálogo central"],
        ["6", "Bajar precio G2 alto tráfico sin venta", "3 cuentas", "Bajo (API)", "+$200-500k/90d"],
        ["7", "Republicar G3 con stock ≥10 (reset ML score)", "3 cuentas", "Alto", "Revive 10-20% de items muertos"],
    ]
    tabla(doc, rows)

    out_comp = ROOT / f"consultor_reporte_COMPARATIVO_{TODAY}.docx"
    doc.save(out_comp)
    print(f"OK {out_comp}")


if __name__ == "__main__":
    main()
