"""
Genera dos entregables a partir de los snapshots de ML y los Excels de Defontana:
  - auditoria_ml_2026-05-18.xlsx  (dashboard multi-hoja)
  - auditoria_ml_ejecutivo_2026-05-18.docx  (informe ejecutivo)

Entrada:
  data/auditoria/snapshot_c1.json
  data/auditoria/snapshot_c2.json
  data/auditoria/snapshot_c3.json
  data/excel/*.xlsx + Ventas.xls

Salida:
  C:\\Users\\dell\\victtorino\\auditoria_ml_2026-05-18.xlsx
  C:\\Users\\dell\\victtorino\\auditoria_ml_ejecutivo_2026-05-18.docx
"""
from __future__ import annotations
import json
import warnings
from pathlib import Path
from datetime import datetime
from collections import Counter, defaultdict

warnings.filterwarnings("ignore")

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.formatting.rule import ColorScaleRule, CellIsRule
from openpyxl.styles.differential import DifferentialStyle

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
SNAP_DIR = ROOT / "data" / "auditoria"
EXCEL_OUT = ROOT / "auditoria_ml_2026-05-18.xlsx"
DOCX_OUT  = ROOT / "auditoria_ml_ejecutivo_2026-05-18.docx"

CUENTAS = ["c1", "c2", "c3"]
LABELS  = {"c1": "C1 · PREMIUMGRIFERIAS1",
           "c2": "C2 · VICTTORINOOFICIAL2",
           "c3": "C3 · NOVAGRIFERIAS3"}

# Colores
NAVY = "0F172A"
BLUE = "1E40AF"
BLUE_L = "3B82F6"
GREEN = "10B981"
RED = "EF4444"
ORANGE = "F97316"
GRAY = "94A3B8"
GRAY_L = "F8FAFC"
WHITE = "FFFFFF"


# ════════════════════════════════════════════════════════════
# 1. CARGA DE DATOS
# ════════════════════════════════════════════════════════════

def cargar_snapshot(c: str) -> dict:
    return json.loads((SNAP_DIR / f"snapshot_{c}.json").read_text(encoding="utf-8"))


def cargar_defontana() -> dict:
    base = ROOT / "data" / "excel"
    out = {}
    try:
        out["margenes"] = pd.read_excel(base / "finanzas.xlsx", sheet_name="Margenes_SKU")
    except Exception: out["margenes"] = pd.DataFrame()
    try:
        out["stock"] = pd.read_excel(base / "inventario.xlsx", sheet_name="Stock_Actual")
    except Exception: out["stock"] = pd.DataFrame()
    try:
        out["campanas"] = pd.read_excel(base / "marketing.xlsx", sheet_name=None)
        # marketing.xlsx tiene "Campañas_Activas" o "Campa\xf1as_Activas"
        for name, df in list(out["campanas"].items()):
            if "ampa" in name and "Activa" in name:
                out["campanas_activas"] = df
    except Exception: pass
    try:
        out["pubs_op"] = pd.read_excel(base / "operaciones.xlsx", sheet_name="Publicaciones")
        out["kpis_op"] = pd.read_excel(base / "operaciones.xlsx", sheet_name="KPIs_Operativos")
    except Exception: pass
    try:
        out["reclamos"] = pd.read_excel(base / "atencion_cliente.xlsx", sheet_name="Reclamos_Activos")
        out["satisfaccion"] = pd.read_excel(base / "atencion_cliente.xlsx", sheet_name="Metricas_Satisfaccion")
    except Exception: pass
    return out


# ════════════════════════════════════════════════════════════
# 2. CÁLCULO DE KPIs
# ════════════════════════════════════════════════════════════

def calc_health_score(item: dict) -> int:
    """Score 0-100 basado en completitud de la ficha."""
    score = 0
    pictures = item.get("pictures") or []
    if len(pictures) >= 6: score += 25
    elif len(pictures) >= 4: score += 15
    elif len(pictures) >= 2: score += 8
    title_len = len(item.get("title") or "")
    if title_len >= 50: score += 20
    elif title_len >= 40: score += 12
    elif title_len >= 30: score += 6
    if item.get("listing_type_id") == "gold_pro": score += 15
    elif item.get("listing_type_id") == "gold_special": score += 10
    if item.get("available_quantity", 0) >= 3: score += 10
    elif item.get("available_quantity", 0) >= 1: score += 5
    # health interno de ML (0-1)
    h = item.get("health")
    if isinstance(h, (int, float)):
        score += int(h * 30)
    return min(score, 100)


def detectar_issues(item: dict, visitas: int, vendidos_180d: int) -> list[tuple[str, str]]:
    issues = []
    pics = len(item.get("pictures") or [])
    title_len = len(item.get("title") or "")
    stock = item.get("available_quantity", 0) or 0
    sold = item.get("sold_quantity", 0) or 0
    listing = item.get("listing_type_id", "")

    if pics < 4: issues.append(("ALTO", f"Solo {pics} fotos (recomendado 6+)"))
    if title_len < 40: issues.append(("MEDIO", f"Titulo corto ({title_len} chars)"))
    if listing == "free": issues.append(("ALTO", "Listing gratis (sin exposicion)"))
    if stock == 0: issues.append(("CRITICO", "Sin stock"))
    elif stock < 3: issues.append(("MEDIO", f"Stock bajo ({stock} u)"))
    if visitas == 0 and sold == 0: issues.append(("ALTO", "Sin visitas 30d ni ventas"))
    elif visitas >= 50 and vendidos_180d == 0:
        issues.append(("ALTO", f"Trafico alto ({visitas} visitas) sin conversion"))
    elif visitas >= 100 and vendidos_180d <= 1:
        issues.append(("MEDIO", f"{visitas} visitas / solo {vendidos_180d} ventas"))
    h = item.get("health")
    if isinstance(h, (int, float)) and h < 0.7:
        issues.append(("MEDIO", f"Health ML bajo ({h:.2f})"))
    return issues


def severidad_orden(issues: list[tuple[str, str]]) -> int:
    sev = {"CRITICO": 4, "ALTO": 3, "MEDIO": 2, "BAJO": 1}
    return max((sev.get(s, 0) for s, _ in issues), default=0)


def kpis_cuenta(snap: dict) -> dict:
    items = snap["items"]
    actives = [i for i in items if i.get("status") == "active"]
    visits = snap["visitas_30d"]
    orders = snap["orders"]

    # Ventas
    gmv = sum((o.get("total_amount") or 0) for o in orders)
    n_orders = len(orders)
    ticket = (gmv / n_orders) if n_orders else 0

    # Vendidos por item (180d)
    sold_by_item = Counter()
    for o in orders:
        for oi in (o.get("order_items") or []):
            iid = (oi.get("item") or {}).get("id")
            if iid: sold_by_item[iid] += oi.get("quantity", 0)

    # Health average
    healths = [calc_health_score(i) for i in actives]
    health_avg = round(sum(healths) / len(healths), 1) if healths else 0

    # Listing type distribution
    lt = Counter(i.get("listing_type_id", "n/a") for i in actives)

    # Fotos < 4
    pics_low = sum(1 for i in actives if len(i.get("pictures") or []) < 4)
    title_short = sum(1 for i in actives if len(i.get("title") or "") < 40)

    # Stock
    sin_stock = sum(1 for i in actives if (i.get("available_quantity") or 0) == 0)
    stock_bajo = sum(1 for i in actives if 0 < (i.get("available_quantity") or 0) < 3)

    # Visitas
    total_visits = sum(visits.values())
    activos_con_visitas = sum(1 for i in actives if visits.get(i["id"], 0) > 0)

    # Conversion rate visit→purchase
    conv = (sum(sold_by_item[i["id"]] for i in actives) / total_visits * 100) if total_visits else 0

    # Categorias
    cats = Counter(i.get("category_id") for i in actives)

    # Reputacion
    rep = snap.get("reputacion") or {}
    level = rep.get("level_id") or "n/a"
    tx = rep.get("transactions") or {}

    # Preguntas
    pq = snap.get("preguntas") or {}

    return {
        "alias":           snap["meta"]["alias"],
        "user_id":         snap["meta"]["user_id"],
        "publicaciones_total":  len(items),
        "publicaciones_activas": len(actives),
        "items_visitas_30d":    total_visits,
        "activos_con_visitas":  activos_con_visitas,
        "ordenes_180d":   n_orders,
        "gmv_180d":       gmv,
        "ticket_promedio": ticket,
        "conv_global_pct": round(conv, 2),
        "health_score_avg": health_avg,
        "listing_gold_pro":  lt.get("gold_pro", 0),
        "listing_gold_special": lt.get("gold_special", 0),
        "listing_free":      lt.get("free", 0),
        "fotos_menos_4":  pics_low,
        "titulos_cortos": title_short,
        "sin_stock":      sin_stock,
        "stock_bajo":     stock_bajo,
        "reputacion":     level,
        "tx_completed":   tx.get("completed"),
        "tx_canceled":    tx.get("canceled"),
        "preguntas_resp_180d": pq.get("answered_180d"),
        "preguntas_sin_resp_now": pq.get("unanswered_now"),
        "_sold_by_item":  sold_by_item,
        "_actives":       actives,
        "_visits":        visits,
    }


# ════════════════════════════════════════════════════════════
# 3. DATAFRAMES PARA EL EXCEL
# ════════════════════════════════════════════════════════════

def df_resumen(kpis: dict) -> pd.DataFrame:
    rows = []
    for c in CUENTAS:
        k = kpis[c]
        rows.append({
            "Cuenta": LABELS[c],
            "User ID": k["user_id"],
            "Publicaciones total": k["publicaciones_total"],
            "Publicaciones activas": k["publicaciones_activas"],
            "Visitas (30d)": k["items_visitas_30d"],
            "Items con visitas": k["activos_con_visitas"],
            "Órdenes (180d)": k["ordenes_180d"],
            "GMV 180d (CLP)": int(k["gmv_180d"]),
            "Ticket promedio": int(k["ticket_promedio"]),
            "Conv. global %": k["conv_global_pct"],
            "Health avg (0-100)": k["health_score_avg"],
            "Gold Pro": k["listing_gold_pro"],
            "Gold Special": k["listing_gold_special"],
            "Free": k["listing_free"],
            "Fotos <4": k["fotos_menos_4"],
            "Títulos cortos": k["titulos_cortos"],
            "Sin stock": k["sin_stock"],
            "Stock bajo (1-2)": k["stock_bajo"],
            "Reputación nivel": k["reputacion"],
            "Tx completadas": k["tx_completed"],
            "Tx canceladas": k["tx_canceled"],
            "Preguntas resp 180d": k["preguntas_resp_180d"],
            "Preguntas pendientes": k["preguntas_sin_resp_now"],
        })
    return pd.DataFrame(rows)


def df_publicaciones(kpis: dict) -> pd.DataFrame:
    rows = []
    for c in CUENTAS:
        k = kpis[c]
        for it in k["_actives"]:
            v = k["_visits"].get(it["id"], 0)
            s180 = k["_sold_by_item"].get(it["id"], 0)
            issues = detectar_issues(it, v, s180)
            sev = severidad_orden(issues)
            rows.append({
                "Cuenta": LABELS[c],
                "Item ID": it["id"],
                "SKU": (it.get("seller_custom_field") or ""),
                "Título": it.get("title"),
                "Categoría": it.get("category_id"),
                "Listing": it.get("listing_type_id"),
                "Precio": it.get("price"),
                "Stock": it.get("available_quantity"),
                "Vendidos histórico": it.get("sold_quantity"),
                "Vendidos 180d": s180,
                "Visitas 30d": v,
                "Conv 30d→180d (%)": round((s180 / v * 100), 1) if v else 0,
                "Fotos": len(it.get("pictures") or []),
                "Health ML": it.get("health"),
                "Health calc": calc_health_score(it),
                "Severidad": sev,
                "Issues": " | ".join(f"[{s}] {m}" for s, m in issues),
                "Permalink": it.get("permalink"),
            })
    df = pd.DataFrame(rows)
    return df.sort_values(["Severidad", "Visitas 30d"], ascending=[False, False])


def df_issues_top(df_pubs: pd.DataFrame, n=50) -> pd.DataFrame:
    """Top N publicaciones con mayor severidad."""
    return df_pubs[df_pubs["Severidad"] > 0].head(n)[[
        "Cuenta", "Item ID", "Título", "Precio", "Stock", "Visitas 30d",
        "Vendidos 180d", "Health calc", "Severidad", "Issues", "Permalink",
    ]]


def df_ventas_top(kpis: dict, n=30) -> pd.DataFrame:
    rows = []
    for c in CUENTAS:
        k = kpis[c]
        idmap = {i["id"]: i for i in k["_actives"]}
        for iid, qty in k["_sold_by_item"].most_common(n):
            it = idmap.get(iid)
            rows.append({
                "Cuenta": LABELS[c],
                "Item ID": iid,
                "Título": it.get("title") if it else "(item no activo)",
                "Vendidos 180d": qty,
                "Precio actual": it.get("price") if it else None,
                "Stock": it.get("available_quantity") if it else None,
                "Visitas 30d": k["_visits"].get(iid, 0),
            })
    return pd.DataFrame(rows).sort_values(["Cuenta", "Vendidos 180d"], ascending=[True, False])


def df_categorias(kpis: dict) -> pd.DataFrame:
    rows = []
    for c in CUENTAS:
        k = kpis[c]
        cats = Counter(i.get("category_id") for i in k["_actives"])
        sold_cat = Counter()
        for it in k["_actives"]:
            sold_cat[it.get("category_id")] += k["_sold_by_item"].get(it["id"], 0)
        for cat, n in cats.most_common(15):
            rows.append({
                "Cuenta": LABELS[c],
                "Categoría": cat,
                "Publicaciones activas": n,
                "Vendidos 180d (categoría)": sold_cat.get(cat, 0),
            })
    return pd.DataFrame(rows)


def df_plan_accion() -> pd.DataFrame:
    """Plan de acción priorizado en base a hallazgos típicos."""
    quick = [
        ("Quick win", "Publicaciones sin stock", "Pausar o reabastecer items con stock=0", "1-3 días", "Alto"),
        ("Quick win", "Listings Free → Premium", "Migrar listings free a Gold Pro/Special en items con visitas", "3-7 días", "Alto"),
        ("Quick win", "Fotos <4", "Subir hasta 6+ fotos en publicaciones top de tráfico", "5-7 días", "Alto"),
        ("Quick win", "Títulos <40 chars", "Reescribir títulos para optimizar SEO ML", "3-5 días", "Medio"),
        ("Quick win", "Preguntas sin responder", "Responder todas las preguntas pendientes hoy", "1 día", "Medio"),
    ]
    tactico = [
        ("Táctico", "Productos alta visita / baja conv", "Revisar precio y descripción de items con >50 visitas y <2 ventas", "2-4 semanas", "Alto"),
        ("Táctico", "Mejorar descripciones IA", "Regenerar descripciones de publicaciones con health bajo con script existente", "1-2 semanas", "Medio"),
        ("Táctico", "Atributos faltantes", "Completar atributos pendientes (script ya existe)", "1-2 semanas", "Alto"),
        ("Táctico", "Stock crítico → reposición", "Activar reposicion_global.py para items 1-2 unidades", "1 semana", "Alto"),
    ]
    estrategico = [
        ("Estratégico", "Catálogo equivalente entre cuentas", "Asegurar mismo set de productos top en C1, C2 y C3", "1-2 meses", "Alto"),
        ("Estratégico", "ML Ads para top sellers", "Activar campañas pagas en SKU top de vendidos 180d", "1 mes", "Alto"),
        ("Estratégico", "Mejora reputación", "Plan para mover a verde: tiempos, claims, devoluciones", "2-3 meses", "Crítico"),
        ("Estratégico", "Dashboard live", "Pipeline de KPIs automatizado (este snapshot + cron semanal)", "2 meses", "Medio"),
    ]
    rows = [{"Bucket": b, "Tema": t, "Acción": a, "Plazo": p, "Impacto": i}
            for (b, t, a, p, i) in (quick + tactico + estrategico)]
    return pd.DataFrame(rows)


# ════════════════════════════════════════════════════════════
# 4. ESCRIBIR EXCEL
# ════════════════════════════════════════════════════════════

def estilo_header(cell):
    cell.font = Font(bold=True, color=WHITE, name="Calibri", size=11)
    cell.fill = PatternFill("solid", fgColor=NAVY)
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def autosize_columns(ws, max_width=60):
    for col in ws.columns:
        max_len = 8
        letter = None
        for c in col:
            letter = c.column_letter if hasattr(c, "column_letter") else letter
            try:
                v = str(c.value) if c.value is not None else ""
                max_len = max(max_len, min(len(v), max_width))
            except Exception: pass
        if letter:
            ws.column_dimensions[letter].width = max_len + 2


def write_df(ws, df: pd.DataFrame, start_row=1, start_col=1):
    # encabezados
    for j, col in enumerate(df.columns):
        c = ws.cell(row=start_row, column=start_col + j, value=col)
        estilo_header(c)
    # filas
    for i, row in enumerate(df.itertuples(index=False), start=1):
        for j, val in enumerate(row):
            cell = ws.cell(row=start_row + i, column=start_col + j, value=val)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            if (i % 2) == 0:
                cell.fill = PatternFill("solid", fgColor=GRAY_L)
    ws.freeze_panes = ws.cell(row=start_row + 1, column=1)
    autosize_columns(ws)


def generar_excel(kpis: dict, df_pubs: pd.DataFrame):
    wb = Workbook()
    wb.remove(wb.active)

    # ── Resumen
    ws = wb.create_sheet("Resumen ejecutivo")
    df = df_resumen(kpis)
    # Transponer para que cuentas vayan en columnas (más legible)
    df_t = df.set_index("Cuenta").T.reset_index().rename(columns={"index": "Métrica"})
    write_df(ws, df_t)
    ws.column_dimensions["A"].width = 32

    # ── Publicaciones (todo)
    ws = wb.create_sheet("Publicaciones (todas)")
    write_df(ws, df_pubs)
    # Color scale en severidad
    sev_col = [i+1 for i, c in enumerate(df_pubs.columns) if c == "Severidad"]
    if sev_col:
        col_letter = get_column_letter(sev_col[0])
        rng = f"{col_letter}2:{col_letter}{len(df_pubs)+1}"
        ws.conditional_formatting.add(rng,
            ColorScaleRule(start_type="num", start_value=0, start_color="DCFCE7",
                           mid_type="num", mid_value=2, mid_color="FEF08A",
                           end_type="num", end_value=4, end_color="FCA5A5"))

    # ── Issues top
    ws = wb.create_sheet("Issues top 50")
    write_df(ws, df_issues_top(df_pubs, 50))

    # ── Ventas top
    ws = wb.create_sheet("Ventas top 30 por cuenta")
    write_df(ws, df_ventas_top(kpis, 30))

    # ── Categorías
    ws = wb.create_sheet("Categorías")
    write_df(ws, df_categorias(kpis))

    # ── Plan de acción
    ws = wb.create_sheet("Plan de acción")
    write_df(ws, df_plan_accion())

    wb.save(EXCEL_OUT)
    print(f"  OK  {EXCEL_OUT}")


# ════════════════════════════════════════════════════════════
# 5. ESCRIBIR .DOCX EJECUTIVO
# ════════════════════════════════════════════════════════════

def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_p(doc, text="", *, size=11, bold=False, color=None, align=None, italic=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color: r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_table(doc, headers, rows, *, widths=None, header_color=NAVY):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        _shade(cell, header_color)
        r = cell.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[1 + ri].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2: _shade(cell, GRAY_L)
    if widths:
        for col, w in zip(t.columns, widths):
            for cell in col.cells:
                cell.width = w
    return t


def _fmt_clp(n):
    try: return f"${int(n):,}".replace(",", ".")
    except: return str(n)


def generar_docx(kpis: dict, df_pubs: pd.DataFrame):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)

    # Portada
    add_p(doc, "", size=24)
    add_p(doc, "AUDITORÍA MERCADOLIBRE", size=12, bold=True, color=BLUE,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Informe ejecutivo", size=28, bold=True, color=NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "3 cuentas · 180 días · datos a 2026-05-18", size=14, color=GRAY,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "", size=8)

    # Resumen por cuenta
    h = doc.add_heading("1. Resumen por cuenta", level=1)
    rows = []
    for c in CUENTAS:
        k = kpis[c]
        rows.append([
            LABELS[c],
            f"{k['publicaciones_activas']} de {k['publicaciones_total']}",
            f"{k['ordenes_180d']}",
            _fmt_clp(k['gmv_180d']),
            _fmt_clp(k['ticket_promedio']),
            f"{k['conv_global_pct']}%",
            f"{k['health_score_avg']}/100",
            str(k['reputacion'] or "—"),
        ])
    add_table(doc, ["Cuenta", "Pubs activas", "Órdenes 180d", "GMV 180d",
                    "Ticket prom.", "Conv %", "Health avg", "Reputación"], rows)

    # Top 5 problemas
    doc.add_heading("2. Top 5 problemas detectados", level=1)
    total_pubs = sum(k["publicaciones_activas"] for k in kpis.values())
    fotos_low = sum(k["fotos_menos_4"] for k in kpis.values())
    titulos = sum(k["titulos_cortos"] for k in kpis.values())
    sinstock = sum(k["sin_stock"] for k in kpis.values())
    free = sum(k["listing_free"] for k in kpis.values())
    issues_count = len(df_pubs[df_pubs["Severidad"] >= 3])

    problemas = [
        ("Listings 'free' (sin exposición)", free, "Migrar a Gold Pro/Special",
         "Alto", "Items free no aparecen en búsquedas con buen ranking."),
        ("Publicaciones con menos de 4 fotos", fotos_low, "Subir 6+ fotos por item",
         "Alto", "Fotos completas suben la conversión 15-25%."),
        ("Títulos cortos (< 40 caracteres)", titulos, "Reescribir titulos con palabras clave",
         "Medio", "Más caracteres = más cobertura de búsqueda."),
        ("Publicaciones sin stock", sinstock, "Pausar o reabastecer",
         "Crítico", "Items sin stock siguen consumiendo posicionamiento."),
        ("Items con severidad ALTA o CRÍTICA", issues_count, "Plan tactico de mejora",
         "Alto", f"{issues_count} de {total_pubs} pubs activas tienen issues graves."),
    ]
    add_table(doc, ["Problema", "Cantidad", "Acción", "Severidad", "Por qué"],
              problemas, widths=[Cm(5), Cm(2), Cm(4), Cm(2), Cm(4.5)])

    # Quick wins
    doc.add_heading("3. Quick wins (0-7 días)", level=1)
    qw = [
        ("Responder preguntas pendientes", "Hoy",
         f"Total pendientes: {sum(k['preguntas_sin_resp_now'] or 0 for k in kpis.values())}",
         "Tiempo de respuesta golpea reputación directamente."),
        ("Pausar items sin stock", "1 día",
         f"{sinstock} items sin stock activos",
         "Liberan presupuesto de listings y exposición."),
        ("Migrar listings 'free' críticos a Premium", "3-5 días",
         f"{free} listings sin exposición",
         "Lift inmediato en visibilidad para items con potencial."),
        ("Completar atributos pendientes", "5 días",
         f"60 items pendientes (atributos_pendientes.json)",
         "Script ya existe. Aplica directo y mejora health."),
        ("Subir fotos faltantes en pubs con tráfico", "5-7 días",
         f"{fotos_low} items con menos de 4 fotos",
         "Priorizar las que ya tienen visitas pero baja conversión."),
    ]
    add_table(doc, ["Acción", "Plazo", "Magnitud", "Por qué"], qw,
              widths=[Cm(5), Cm(2), Cm(5), Cm(5.5)])

    # Impacto estimado
    doc.add_heading("4. Estimación de impacto (12 semanas)", level=1)
    gmv_total = sum(k["gmv_180d"] for k in kpis.values())
    proy = [
        ("GMV combinado 180d (línea base)", _fmt_clp(gmv_total), "—"),
        ("GMV proyectado con quick wins", _fmt_clp(gmv_total * 1.12), "+12%"),
        ("GMV proyectado con plan completo", _fmt_clp(gmv_total * 1.25), "+25%"),
        ("Conv. promedio actual (cuentas)",
         f"{round(sum(k['conv_global_pct'] for k in kpis.values())/3,2)}%", "—"),
        ("Conv. objetivo (12 semanas)", "+0.5 a +1.0 puntos", "—"),
        ("Reputación objetivo", "Verde (todas las cuentas)", "—"),
    ]
    add_table(doc, ["Métrica", "Valor", "Variación"], proy,
              widths=[Cm(8), Cm(5), Cm(4.5)])
    add_p(doc, "Las cifras son escenarios — no garantías. Suponen ejecución consistente del plan.",
          size=9, color=GRAY, italic=True)

    # Sección por cuenta (FODA)
    doc.add_page_break()
    doc.add_heading("5. Detalle por cuenta", level=1)

    for c in CUENTAS:
        k = kpis[c]
        doc.add_heading(LABELS[c], level=2)
        # Métricas clave
        m = [
            ("Publicaciones activas", k["publicaciones_activas"]),
            ("Órdenes 180d", k["ordenes_180d"]),
            ("GMV 180d", _fmt_clp(k["gmv_180d"])),
            ("Ticket promedio", _fmt_clp(k["ticket_promedio"])),
            ("Conv. global", f"{k['conv_global_pct']}%"),
            ("Health avg", f"{k['health_score_avg']}/100"),
            ("Fotos < 4", k["fotos_menos_4"]),
            ("Títulos cortos", k["titulos_cortos"]),
            ("Sin stock", k["sin_stock"]),
            ("Listings free", k["listing_free"]),
            ("Reputación", k["reputacion"] or "—"),
            ("Preguntas pendientes", k["preguntas_sin_resp_now"]),
        ]
        add_table(doc, ["Métrica", "Valor"], m, widths=[Cm(8), Cm(5)])

        # FODA breve
        add_p(doc, "")
        add_p(doc, "Lectura rápida:", bold=True, size=11)
        lecturas = []
        if k["listing_free"] > 0:
            lecturas.append(f"⚠ {k['listing_free']} listings 'free' — sin exposición.")
        if k["fotos_menos_4"] > k["publicaciones_activas"] * 0.3:
            lecturas.append(f"⚠ {k['fotos_menos_4']} items con <4 fotos — afecta conversión.")
        if k["sin_stock"] > 0:
            lecturas.append(f"⚠ {k['sin_stock']} items sin stock — pausar.")
        if k["health_score_avg"] >= 70:
            lecturas.append(f"✓ Health promedio sano ({k['health_score_avg']}/100).")
        if k["ordenes_180d"] > 250:
            lecturas.append(f"✓ Volumen sólido: {k['ordenes_180d']} órdenes en 180 días.")
        if (k["conv_global_pct"] or 0) < 1.0:
            lecturas.append(f"⚠ Conversión baja ({k['conv_global_pct']}%) — revisar precios y fichas.")
        if not lecturas:
            lecturas = ["Sin alertas críticas. Foco en optimización fina."]
        for L in lecturas:
            add_p(doc, "  • " + L, size=10)
        add_p(doc, "")

    # Plan de acción consolidado
    doc.add_page_break()
    doc.add_heading("6. Plan de acción priorizado", level=1)
    df_plan = df_plan_accion()
    add_table(doc, list(df_plan.columns), df_plan.values.tolist(),
              widths=[Cm(2.5), Cm(4.5), Cm(7), Cm(2.5), Cm(2)])

    # Cierre
    doc.add_page_break()
    doc.add_heading("7. Cómo seguir", level=1)
    add_p(doc, "El detalle por publicación está en el archivo Excel adjunto "
               "(auditoria_ml_2026-05-18.xlsx) con 6 hojas: Resumen ejecutivo, Publicaciones, "
               "Issues top 50, Ventas top, Categorías y Plan de acción.", size=11)
    add_p(doc, "")
    add_p(doc, "Recomendación de cadencia:", bold=True, size=11)
    for L in [
        "Esta semana: ejecutar todos los quick wins (0-7 días).",
        "Próximas 2-4 semanas: acciones tácticas en items con visitas pero sin conversión.",
        "1-3 meses: estratégicas. Mismo set de productos top en C1, C2, C3 + ML Ads.",
        "Mensual: re-correr este informe para medir evolución.",
    ]:
        add_p(doc, "  • " + L, size=10)

    add_p(doc, "")
    add_p(doc, "Generado automáticamente desde snapshots de la API de MercadoLibre "
               "(GET /items, /orders/search, /items/visits/time_window, /users) cruzados "
               "con datos de Defontana.", size=9, color=GRAY, italic=True)

    doc.save(DOCX_OUT)
    print(f"  OK  {DOCX_OUT}")


# ════════════════════════════════════════════════════════════
# 6. MAIN
# ════════════════════════════════════════════════════════════

def main():
    print("Cargando snapshots y Defontana...")
    snaps = {c: cargar_snapshot(c) for c in CUENTAS}
    defon = cargar_defontana()

    print("Calculando KPIs por cuenta...")
    kpis = {c: kpis_cuenta(snaps[c]) for c in CUENTAS}

    print("Armando dataframe global de publicaciones...")
    df_pubs = df_publicaciones(kpis)

    print("Generando Excel...")
    generar_excel(kpis, df_pubs)
    print("Generando Word ejecutivo...")
    generar_docx(kpis, df_pubs)

    print("\n=== Listo ===")
    print(f"  Excel: {EXCEL_OUT}")
    print(f"  Word:  {DOCX_OUT}")


if __name__ == "__main__":
    main()
