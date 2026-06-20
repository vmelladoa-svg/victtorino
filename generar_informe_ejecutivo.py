"""
Informe ejecutivo DOCX a partir de analisis.pkl + snapshots refrescados.
Genera auditoria_ml_ejecutivo_<fecha>.docx
"""
import pickle
from datetime import date, datetime, timedelta, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
ANALISIS = ROOT / "data" / "auditoria" / "analisis.pkl"
TODAY = date.today().isoformat()
OUT = ROOT / f"auditoria_ml_ejecutivo_{TODAY}.docx"

BRAND = RGBColor(0x1F, 0x4E, 0x78)
WARN = RGBColor(0xC0, 0x39, 0x2B)
OK = RGBColor(0x16, 0x7A, 0x3F)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=BRAND):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = "Calibri"
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p


def add_bullet(doc, items, bold_lead=False):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if bold_lead and ": " in it:
            lead, rest = it.split(": ", 1)
            r = p.add_run(lead + ": "); r.bold = True; r.font.size = Pt(11)
            r2 = p.add_run(rest); r2.font.size = Pt(11)
        else:
            r = p.add_run(it); r.font.size = Pt(11)


def calc_conv30(snapshots):
    cutoff = datetime.now(timezone.utc) - timedelta(days=30)
    out = {}
    for c, snap in snapshots.items():
        o30 = [o for o in snap["orders"]
               if datetime.fromisoformat(o["date_created"].replace("Z","+00:00")) >= cutoff]
        v30 = sum(snap.get("visitas_30d",{}).values())
        gmv30 = sum(o.get("total_amount",0) for o in o30)
        out[c] = {
            "orders30": len(o30),
            "gmv30": int(gmv30),
            "aov30": int(gmv30/len(o30)) if o30 else 0,
            "conv30": round(len(o30)/v30*100, 2) if v30 else 0,
            "visits30": v30,
        }
    return out


def add_kpi_table(doc, df_res, conv30):
    by = {r["Cuenta"]: r for _, r in df_res.iterrows()}
    rows = [("KPI", "C1 PREMIUMGRIFERIAS1", "C2 VICTTORINOOFICIAL2", "C3 NOVAGRIFERIAS3", "Total / Mejor")]

    def line(label, vals, total_calc=None, lower_better=False, suffix="", fmt_thousands=False):
        # vals: tuple of 3 strings or numbers
        all_num = all(isinstance(v,(int,float)) for v in vals)
        if total_calc == "sum" and all_num:
            total = sum(vals)
        elif total_calc == "avg" and all_num:
            total = round(sum(vals)/len(vals), 2)
        else:
            total = ""
        fmt_v = []
        if all_num:
            best = min(vals) if lower_better else max(vals)
            for v in vals:
                if fmt_thousands and abs(v) >= 1000:
                    s = f"{v:,.0f}"
                elif isinstance(v,float):
                    s = f"{v:.2f}"
                else:
                    s = str(v)
                s = s + suffix
                if v == best:
                    s = "★ " + s
                fmt_v.append(s)
            if isinstance(total,(int,float)):
                total_s = f"{total:,.0f}{suffix}" if fmt_thousands else f"{total}{suffix}"
            else:
                total_s = ""
        else:
            fmt_v = [str(v)+suffix for v in vals]
            total_s = ""
        return [label] + fmt_v + [total_s]

    rows.append(line("Publicaciones activas", (by["C1"]["Activas"], by["C2"]["Activas"], by["C3"]["Activas"]), "sum"))
    rows.append(line("Visitas 30 días", (by["C1"]["Visitas30d"], by["C2"]["Visitas30d"], by["C3"]["Visitas30d"]), "sum", fmt_thousands=True))
    rows.append(line("Items con tráfico", (by["C1"]["ItemsConTrafico"], by["C2"]["ItemsConTrafico"], by["C3"]["ItemsConTrafico"]), "sum"))
    rows.append(line("Órdenes 180 días", (by["C1"]["Ordenes180d"], by["C2"]["Ordenes180d"], by["C3"]["Ordenes180d"]), "sum"))
    rows.append(line("GMV 180 días (CLP)", (by["C1"]["GMV180d"], by["C2"]["GMV180d"], by["C3"]["GMV180d"]), "sum", suffix=" $", fmt_thousands=True))
    rows.append(line("Órdenes 30 días", (conv30["C1"]["orders30"], conv30["C2"]["orders30"], conv30["C3"]["orders30"]), "sum"))
    rows.append(line("GMV 30 días (CLP)", (conv30["C1"]["gmv30"], conv30["C2"]["gmv30"], conv30["C3"]["gmv30"]), "sum", suffix=" $", fmt_thousands=True))
    rows.append(line("AOV 30 días (CLP)", (conv30["C1"]["aov30"], conv30["C2"]["aov30"], conv30["C3"]["aov30"]), "avg", suffix=" $", fmt_thousands=True))
    rows.append(line("Conv 30d (ord/vis) %", (conv30["C1"]["conv30"], conv30["C2"]["conv30"], conv30["C3"]["conv30"]), "avg", suffix="%"))
    rows.append(line("Health avg /100", (by["C1"]["HealthAvg"], by["C2"]["HealthAvg"], by["C3"]["HealthAvg"]), "avg"))
    rows.append(line("Items <4 fotos", (by["C1"]["FotosMenos4"], by["C2"]["FotosMenos4"], by["C3"]["FotosMenos4"]), "sum", lower_better=True))
    rows.append(line("Stock crítico (1-2 u)", (by["C1"]["StockCritico<3"], by["C2"]["StockCritico<3"], by["C3"]["StockCritico<3"]), "sum", lower_better=True))
    rows.append(line("Reputación", (str(by["C1"]["RepNivel"]), str(by["C2"]["RepNivel"]), str(by["C3"]["RepNivel"]))))
    rows.append(line("Tx canceladas", (by["C1"]["TxCanceladas"], by["C2"]["TxCanceladas"], by["C3"]["TxCanceladas"]), "sum", lower_better=True))
    rows.append(line("Preguntas pendientes", (by["C1"]["PreguntasPendientes"], by["C2"]["PreguntasPendientes"], by["C3"]["PreguntasPendientes"]), "sum", lower_better=True))

    table = doc.add_table(rows=len(rows), cols=5)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if ri == 0:
                shade_cell(cell, "1F4E78")
    return table


def build_foda(cuenta, by, conv30, df_pub):
    r = by[cuenta]; c30 = conv30[cuenta]
    fortalezas = []; debilidades = []; oportunidades = []; amenazas = []

    if str(r["RepNivel"]) == "5_green":
        fortalezas.append(f"Reputación verde nivel 5 ({r['TxCompletadas']} tx completadas, {r['TxCanceladas']} canceladas)")
    if r["SinStock"] == 0:
        fortalezas.append("Sin publicaciones quebradas (stock 0) en activas")
    if r["ListingsFree"] == 0:
        fortalezas.append("Catálogo 100% en Gold (Pro o Special) — sin Free")
    if c30["conv30"] >= 3:
        fortalezas.append(f"Conv 30d {c30['conv30']}% sobre referencia ML (~2.5-3%)")
    if r["HealthAvg"] >= 50:
        fortalezas.append(f"Health avg {r['HealthAvg']}/100 — relativamente sano")
    if r["PreguntasPendientes"] == 0:
        fortalezas.append("Sin preguntas pendientes (atención al cliente al día)")

    if r["FotosMenos4"] > 30:
        debilidades.append(f"{r['FotosMenos4']} publicaciones con <4 fotos (~{r['FotosMenos4']/r['Activas']*100:.0f}% del catálogo) — pierden conversión")
    if r["StockCritico<3"] > 30:
        debilidades.append(f"{r['StockCritico<3']} ítems con stock crítico (1-2 u) — riesgo OOS inminente")
    if r["HealthAvg"] < 50:
        debilidades.append(f"Health avg {r['HealthAvg']}/100 bajo — atributos y ficha incompletos")
    no_traf = r["Activas"] - r["ItemsConTrafico"]
    if no_traf >= 30:
        debilidades.append(f"{no_traf} activas sin visitas en 30d (~{no_traf/r['Activas']*100:.0f}% del catálogo) — SEO interno débil")
    if r["TitulosCortos<40"] > 15:
        debilidades.append(f"{r['TitulosCortos<40']} títulos cortos (<40 chars) — desperdician keywords")
    if c30["conv30"] < 3:
        debilidades.append(f"Conv 30d {c30['conv30']}% bajo referencia (~3%)")

    tns = df_pub[(df_pub["Cuenta"]==cuenta) & (df_pub["Visitas30d"]>=50) & (df_pub["Vendidos180d"]==0)]
    if len(tns) > 0:
        oportunidades.append(f"{len(tns)} ítems con ≥50 visitas y 0 ventas — revisar precio/oferta/copy")
    qw = df_pub[(df_pub["Cuenta"]==cuenta) & (df_pub["Visitas30d"]>=10) & (df_pub["HealthCalc"]<70) & (df_pub["Vendidos180d"]>0)]
    if len(qw) > 0:
        rev_pot = int(qw["Revenue180d"].sum() * 0.3)
        oportunidades.append(f"{len(qw)} quick wins con +30% revenue potencial = ${rev_pot:,} CLP/180d adicional")
    if cuenta != "C3":
        oportunidades.append("Replicar fichas y categorización de C3 (lidera en visitas + health)")
    oportunidades.append("Mercado Ads sin activar — espacio para top-of-funnel pagado (presupuesto $10k/día ya configurado en C3)")
    oportunidades.append("Agregar video a top 30 SKUs (+10-15% conversión típica en ML)")

    if r["TxCanceladas"] and r["TxCompletadas"] and (r["TxCanceladas"] / r["TxCompletadas"]) > 0.08:
        amenazas.append(f"Cancelaciones {r['TxCanceladas']}/{r['TxCompletadas']} ({r['TxCanceladas']/r['TxCompletadas']*100:.1f}%) — riesgo bajar nivel")
    if r["MetricaDelayedRate%"] and r["MetricaDelayedRate%"] > 0.015:
        amenazas.append(f"Delayed handling {r['MetricaDelayedRate%']*100:.2f}% sobre umbral ML")
    amenazas.append("Sin cruce sistemático ML ↔ Defontana (riesgo OOS y costo desactualizado)")
    amenazas.append("Competencia agresiva categoría baño/cocina — sin Ads se pierde share")

    return fortalezas, debilidades, oportunidades, amenazas


def main():
    data = pickle.loads(ANALISIS.read_bytes())
    df_res = data["df_resumen"]
    df_pub = data["df_publicaciones"]
    snapshots = data["snapshots"]
    by = {r["Cuenta"]: r for _, r in df_res.iterrows()}
    conv30 = calc_conv30(snapshots)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = title.add_run("AUDITORÍA INTEGRAL\nMERCADOLIBRE — 3 CUENTAS")
    r0.bold = True; r0.font.size = Pt(28); r0.font.color.rgb = BRAND
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = sub.add_run("PREMIUMGRIFERIAS1 · VICTTORINOOFICIAL2 · NOVAGRIFERIAS3")
    r1.italic = True; r1.font.size = Pt(14)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = meta.add_run(f"\nFecha: {TODAY}    |    {int(df_res['Activas'].sum())} activas    |    GMV 180d ${int(df_res['GMV180d'].sum()):,} CLP")
    r2.font.size = Pt(11)

    doc.add_paragraph()
    add_heading(doc, "Resumen Ejecutivo", level=1)
    total_gmv = int(df_res['GMV180d'].sum())
    total_o = int(df_res['Ordenes180d'].sum())
    total_v = int(df_res['Visitas30d'].sum())
    total_act = int(df_res['Activas'].sum())
    sin_v = int((df_pub["Visitas30d"]==0).sum())
    pocas_f = int((df_pub["Fotos"]<4).sum())
    crit = int(df_pub["Issues"].str.contains("CRITICO", na=False).sum())
    qw_total = int(((df_pub["Visitas30d"]>=10) & (df_pub["HealthCalc"]<70) & (df_pub["Vendidos180d"]>0)).sum())

    add_para(doc,
        f"Las 3 cuentas suman {total_act} publicaciones activas, "
        f"${total_gmv:,} CLP de GMV en 180 días ({total_o} órdenes) "
        f"y {total_v:,} visitas en los últimos 30 días. "
        f"La conversión 30d real (órdenes/visitas) por cuenta es "
        f"C1 {conv30['C1']['conv30']}%, C2 {conv30['C2']['conv30']}%, C3 {conv30['C3']['conv30']}%, "
        f"con AOV C1 ${conv30['C1']['aov30']:,}, C2 ${conv30['C2']['aov30']:,}, C3 ${conv30['C3']['aov30']:,}. "
        f"C3 lidera en volumen y tráfico (3.211 visitas, $4.1M en 30d) y C2 monetiza mejor cada visita "
        f"(4.38% conv vs 2.72% en C1). "
        f"El health avg del catálogo es 48.8/100 — bajo para benchmarks ML (objetivo >65)."
    )
    add_para(doc, "Top 5 problemas que están frenando ventas:", bold=True, after=2)
    add_bullet(doc, [
        f"Calidad de ficha pobre: {pocas_f} publicaciones con <4 fotos (~{pocas_f/total_act*100:.0f}% del catálogo). Impacto directo en conversión.",
        f"Catálogo poco descubrible: {sin_v} ítems sin una sola visita en 30d (~{sin_v/total_act*100:.0f}%). SEO de título/categorías/atributos por optimizar.",
        f"Sin publicidad activa: 0 campañas Mercado Ads detectadas en las 3 cuentas (C3 tiene presupuesto $10k/día configurado pero apagado).",
        f"Inconsistencia entre cuentas: C3 captura 3× el tráfico de C1; mismo rubro y stock similar. Replicación de fichas top de C3 a C1/C2 es win seguro.",
        f"Tiempos de despacho marginales: 3 métricas en C1 sobre umbral ML (delayed_handling, claims, cancellations) — riesgo de bajar de nivel verde.",
    ])

    add_para(doc, "Recomendaciones prioritarias (quick wins 0-7 días):", bold=True, after=2)
    add_bullet(doc, [
        f"Subir fotos faltantes en los {qw_total} ítems quick-win (tráfico ≥10, health <70, ya venden): proyecta +30% revenue en ese subset.",
        f"Atender {crit} publicaciones críticas (tráfico sin venta, precio cero, stock 0).",
        "Activar campaña Mercado Ads inicial en C3 (top 10 bestsellers) con presupuesto $10k/día existente — ROAS típico 5-8× en baño/cocina.",
        "Verificar fix de procesing_time en Seller Center C1 (pendiente histórico — acción manual por Victor).",
    ])

    add_para(doc, "Estimación de impacto a 90 días si se ejecuta el plan:", bold=True, after=2)
    add_bullet(doc, [
        f"GMV 90d (proyección base, mitad de los 180d): ${int(total_gmv/2):,}. Meta conservadora (+15%): ${int(total_gmv/2*1.15):,}. Meta agresiva (+40%): ${int(total_gmv/2*1.4):,}.",
        "Health avg subir de 48.8 a 65+ con script de atributos+fotos.",
        "Conv 30d objetivo: 5% en las 3 cuentas (hoy 2.7-4.4%).",
    ])

    doc.add_page_break()
    add_heading(doc, "KPIs comparativa 3 cuentas", level=1)
    add_kpi_table(doc, df_res, conv30)

    for cuenta in ("C1","C2","C3"):
        doc.add_page_break()
        r = by[cuenta]
        c30 = conv30[cuenta]
        add_heading(doc, f"{cuenta} — {r['Alias']}", level=1)
        add_para(doc,
            f"User ID {r['UserID']} · Reputación {r['RepNivel']} ({r['PowerSeller']}) · "
            f"{r['Activas']} activas · {r['Ordenes180d']} órdenes 180d · "
            f"GMV ${r['GMV180d']:,} CLP",
            italic=True, size=10)

        add_heading(doc, "Performance 30 días (apples-to-apples)", level=2)
        add_bullet(doc, [
            f"Visitas: {c30['visits30']:,}",
            f"Órdenes: {c30['orders30']}",
            f"GMV: ${c30['gmv30']:,} CLP",
            f"AOV: ${c30['aov30']:,}",
            f"Conv 30d: {c30['conv30']}%",
        ])

        f, d, o, a = build_foda(cuenta, by, conv30, df_pub)
        add_heading(doc, "Análisis FODA", level=2)
        for label, items, color in (("Fortalezas",f,OK),("Debilidades",d,WARN),
                                     ("Oportunidades",o,BRAND),("Amenazas",a,WARN)):
            p = doc.add_paragraph()
            run = p.add_run(label + ":")
            run.bold = True; run.font.size = Pt(11); run.font.color.rgb = color
            add_bullet(doc, items)

        add_heading(doc, f"Top 5 publicaciones a intervenir en {cuenta}", level=2)
        sub = df_pub[df_pub["Cuenta"]==cuenta].sort_values(["Severidad","Visitas30d"], ascending=[False,False]).head(5)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Item","Título","Visitas/Ventas","Health","Acción sugerida"]):
            hdr[i].text = h
            shade_cell(hdr[i], "1F4E78")
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True; run.font.color.rgb = RGBColor(255,255,255); run.font.size = Pt(10)
        for _, p in sub.iterrows():
            cells = tbl.add_row().cells
            cells[0].text = str(p["ItemID"])
            cells[1].text = (str(p["Título"]) or "")[:55]
            cells[2].text = f"{p['Visitas30d']}/{p['Vendidos180d']}"
            cells[3].text = str(p["HealthCalc"])
            cells[4].text = (str(p["Acciones"]) or "")[:80]
            for c in cells:
                for para in c.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        # Top vendedores
        add_heading(doc, f"Top 5 bestsellers de {cuenta}", level=2)
        tops = df_pub[df_pub["Cuenta"]==cuenta].nlargest(5, "Vendidos180d")
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Item","Título","Stock","Vendidos 180d","Revenue 180d"]):
            hdr[i].text = h
            shade_cell(hdr[i], "1F4E78")
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True; run.font.color.rgb = RGBColor(255,255,255); run.font.size = Pt(10)
        for _, p in tops.iterrows():
            cells = tbl.add_row().cells
            cells[0].text = str(p["ItemID"])
            cells[1].text = (str(p["Título"]) or "")[:55]
            cells[2].text = str(p["Stock"])
            cells[3].text = str(p["Vendidos180d"])
            cells[4].text = f"${int(p['Revenue180d']):,}"
            for c in cells:
                for para in c.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    doc.add_page_break()
    add_heading(doc, "Plan de Acción Integrado", level=1)

    add_heading(doc, "QUICK WINS (0-7 días)", level=2, color=OK)
    add_bullet(doc, [
        "Sin stock: 0 ítems hoy en las 3 cuentas (mantener vigilancia weekly).",
        f"Fotos <4: subir fotos en los 161 ítems afectados, empezar por los {qw_total} quick-wins con tráfico.",
        "Preguntas pendientes: 1 ítem en C2 — responder hoy.",
        "Tráfico sin venta (≥50 visitas, 0 ventas 180d): revisar 7 ítems — precio vs competencia, fotos, oferta.",
        "Mercado Ads en C3: activar campaña inicial con top 10 SKUs y presupuesto $10k/día existente.",
        "Verificar fix procesing_time en Seller Center C1 (acción manual por Victor).",
    ], bold_lead=True)

    add_heading(doc, "MEJORAS TÁCTICAS (1-4 semanas)", level=2, color=BRAND)
    add_bullet(doc, [
        "Replicar catálogo C3→C1/C2: replicar fichas, atributos y categorías de los 30 bestsellers de C3 en las otras 2 cuentas.",
        "Atributos pendientes: completar con script actualizar_atributos.py (afecta ~80% del catálogo con health <60).",
        "Optimizar SEO de los 150 ítems sin visitas 30d (título palabra clave, categoría, foto primaria).",
        "Activar Mercado Ads en C1 y C2 una vez validado ROI en C3 (presupuesto inicial $5k/día por cuenta).",
        "Pipeline de KPIs semanal: cron del audit cada lunes + Excel automatizado.",
        f"Reducir cancelaciones C1 ({by['C1']['TxCanceladas']}/{by['C1']['TxCompletadas']}) — diagnóstico motivos top y SOP.",
    ], bold_lead=True)

    add_heading(doc, "CAMBIOS ESTRATÉGICOS (1-3 meses)", level=2, color=BRAND)
    add_bullet(doc, [
        "Equivalencia de catálogo entre cuentas: mismo top 50 bestsellers en C1, C2 y C3.",
        "Video por SKU bestseller: grabar 30 videos en los primeros 60 días.",
        "Cruce ML ↔ Defontana real: integrar API Token Defontana para tener costo real, stock live y margen exacto por publicación (hoy solo 15 SKUs en muestra).",
        "Programa de compra repetida: cupón o descuento para buyers ML que ya compraron.",
        "A/B test de precio en categorías clave (lavaplatos, llaves) con tracking 30d.",
        "Dashboard live (Streamlit/Looker) con snapshot diario.",
    ], bold_lead=True)

    add_heading(doc, "Proyección de impacto a 90 días", level=2)
    add_para(doc, "Escenarios construidos sobre base GMV 90d ≈ $13.0M (la mitad de los 180d).", italic=True)
    proj_table = doc.add_table(rows=4, cols=4)
    proj_table.style = "Light Grid Accent 1"
    hdr = ["Escenario","Acciones ejecutadas","GMV 90d proyectado","Δ vs base"]
    for i, h in enumerate(hdr):
        proj_table.rows[0].cells[i].text = h
        shade_cell(proj_table.rows[0].cells[i], "1F4E78")
        for run in proj_table.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True; run.font.color.rgb = RGBColor(255,255,255)
    base90 = int(total_gmv/2)
    esc = [
        ("Sin acción","Mantener estado actual",f"${base90:,}","0%"),
        ("Conservador","Quick wins (fotos, sin stock, preguntas, ads C3)",f"${int(base90*1.15):,}","+15%"),
        ("Agresivo","Quick wins + tácticos (atributos, ads 3 cuentas, replicación C3→C1/C2)",f"${int(base90*1.40):,}","+40%"),
    ]
    for i, row in enumerate(esc, 1):
        for j, v in enumerate(row):
            proj_table.rows[i].cells[j].text = v
            for run in proj_table.rows[i].cells[j].paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_page_break()
    add_heading(doc, "Anexos y limitaciones", level=1)
    add_para(doc, "Esta auditoría se acompaña de:", bold=True)
    add_bullet(doc, [
        f"auditoria_ml_{TODAY}.xlsx — dashboard interactivo (10 hojas)",
        f"data/auditoria/snapshot_c1.json | snapshot_c2.json | snapshot_c3.json — snapshots crudos ML API",
        f"data/auditoria/analisis.pkl — DataFrames consolidados",
    ])
    add_para(doc, "Limitaciones conocidas:", bold=True)
    add_bullet(doc, [
        "Mercado Ads API v1 retorna 404 (deprecada). La evaluación de campañas se hizo con scrape previo + UI cache (C3 confirma 0 campañas activas). Para automatizar, integrar Mercado Ads API v2 cuando ML habilite acceso a las cuentas.",
        "Defontana: solo 15 SKUs en la muestra Excel — no permite cruce sistemático con los ~530 ítems ML. Resolver con API Token Defontana (pendiente).",
        "El cálculo de 'ConvRate%' por publicación usa ventas 180d / visitas 30d — proxy direccional, no estricto. La conv 30d real está en la tabla KPI (apples-to-apples).",
        "Health ML (campo nativo) está poblado solo en ~30-40% de los ítems. El 'HealthCalc' del dashboard es nuestro proxy 0-100 sobre fotos/título/listing/tráfico/conversión.",
    ])

    doc.save(OUT)
    print(f"OK Informe: {OUT}  ({OUT.stat().st_size//1024} KB)")


if __name__ == "__main__":
    main()
