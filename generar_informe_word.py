"""
Genera la version WORD del informe ejecutivo para revisar manualmente.
Equivalente al informe_revision_<fecha>.xlsx pero en formato Word imprimible.
"""
from __future__ import annotations
import sys
import time
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import load_workbook

ROOT = Path(__file__).parent
PLAN = ROOT / f"plan_accion_{time.strftime('%Y-%m-%d')}.xlsx"
DIAG = ROOT / f"margen_diagnostico_{time.strftime('%Y-%m-%d')}.xlsx"
DST = ROOT / f"informe_revision_{time.strftime('%Y-%m-%d')}.docx"

IVA = 0.19
ML_CHANNEL_FEE = 0.32
NETO_FACTOR = (1 / (1 + IVA)) * (1 - ML_CHANNEL_FEE)
VENTAS_GANA = 3.0
VENTAS_PIERDE = 0.3


def cargar(path: Path, hoja: str) -> list[dict]:
    wb = load_workbook(path, read_only=True)
    if hoja not in wb.sheetnames:
        return []
    ws = wb[hoja]
    rows = ws.iter_rows(values_only=True)
    h = next(rows)
    return [dict(zip(h, r)) for r in rows]


def calc_margen(price, costo):
    if not price or not costo or price <= 0 or costo <= 0:
        return 0
    return (price / (1 + IVA)) * (1 - ML_CHANNEL_FEE) - costo


def fmt(v):
    try:
        return f"${int(v):,}".replace(",", ".")
    except Exception:
        return "-"


def shade_cell(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color="1A237E"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_table_simple(doc, headers, rows, header_color="455A64", header_font_color="FFFFFF", widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor.from_string(header_font_color)
                r.font.size = Pt(10)
    for row_data in rows:
        rc = t.add_row().cells
        for i, val in enumerate(row_data):
            rc[i].text = str(val) if val is not None else ""
            for p in rc[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return t


def main() -> int:
    if not PLAN.exists():
        print(f"ERROR: no existe {PLAN.name}")
        return 1

    acciones_si = [a for a in cargar(PLAN, "ACCIONES") if a.get("aprobado") == "SI"]
    todos = cargar(DIAG, "TODOS_LOS_LISTINGS") if DIAG.exists() else []
    todas_acciones = cargar(PLAN, "ACCIONES")

    from collections import Counter
    accion_cuenta = Counter((a["accion"], a["cuenta"]) for a in acciones_si)
    total_por_cuenta = Counter(a["cuenta"] for a in acciones_si)

    # proyeccion totales
    total_margen_actual = 0
    total_margen_proy = 0
    total_margen_transf = 0
    for a in acciones_si:
        m_act_unidad = calc_margen(a.get("precio_actual"), a.get("costo"))
        m_proy_unidad = calc_margen(a.get("precio_sugerido") or a.get("precio_actual"),
                                     a.get("costo"))
        if a["accion"] == "PAUSAR":
            total_margen_actual += VENTAS_PIERDE * m_act_unidad
            total_margen_proy += 0
            total_margen_transf += VENTAS_PIERDE * 0.8 * m_act_unidad
        else:
            total_margen_actual += VENTAS_PIERDE * m_act_unidad
            total_margen_proy += VENTAS_GANA * m_proy_unidad
    impacto_mes = total_margen_proy + total_margen_transf - total_margen_actual

    # =====================================================
    # DOCUMENT
    # =====================================================
    doc = Document()

    # margenes y page
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Portada / Header
    titulo = doc.add_heading("INFORME EJECUTIVO", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = add_para(doc, "Optimizacion Catalogo MercadoLibre — Plan de Accion",
                   bold=True, size=14, color="1A237E", align="center")
    add_para(doc, f"Generado: {time.strftime('%Y-%m-%d %H:%M')}  |  Cuentas: C1, C2, C3  |  533 publicaciones analizadas",
             size=10, color="64748B", align="center")
    doc.add_paragraph()

    # ============== 1. DIAGNOSTICO ==============
    add_heading(doc, "1. Diagnostico actual", level=1)
    diag_lines = [
        "533 publicaciones en catalog listing repartidas entre las 3 cuentas (C1: 170, C2: 176, C3: 187).",
        "152 fichas de catalogo tienen overlap entre cuentas tuyas (canibalizacion interna).",
        "De esas, 40 tienen al menos una cuenta con stock real — donde la canibalizacion duele de verdad.",
        "51 publicaciones pierden buy box vs sellers externos (FACTORYNETCL, VIBRA TOOLS, OUTLETSEVENCL concentran 41%).",
        "12 de 13 'lookalikes sospechosos' resultaron ser EL MISMO producto vendido por distribuidores tuyos (problema de control de canal, no marketing).",
        "13 items vendiendose a perdida (revision urgente — probable error de carga de precio o costo desactualizado).",
        "Costos Defontana (netos) cruzaron al 95.3% con items ML via SELLER_SKU.",
    ]
    for line in diag_lines:
        p = doc.add_paragraph(line, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(10.5)
    doc.add_paragraph()

    # ============== 2. ESTRATEGIA ==============
    add_heading(doc, "2. Estrategia que implementamos", level=1)
    estr = [
        ("Principio 1 — Una cuenta nuestra por ficha",
         "En cada catalog_product_id donde tenemos overlap interno, dejamos solo a la cuenta que YA gana el buy box. Las demas se pausan. NO es 'siempre C3': es 'la mas eficiente actual'."),
        ("Principio 2 — Bajamos precio solo si el margen aguanta",
         "Aplicamos BAJAR_PRECIO solo cuando el margen final queda >=15% Y la caida es <=10 puntos. Descarta los casos donde el externo es estructuralmente mas barato."),
        ("Principio 3 — No tocamos lo que funciona",
         "429 fichas (84% del catalogo) no se tocan. Ya ganamos o no hay competencia. Margenes muy altos."),
        ("Principio 4 — Decisiones comerciales separadas",
         "Control de canal/distribuidores, salir del catalogo, brand protection: requieren input tuyo. No las decide este plan."),
    ]
    for nombre, texto in estr:
        p = doc.add_paragraph()
        r = p.add_run(nombre + ":  ")
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor.from_string("1A237E")
        r2 = p.add_run(texto)
        r2.font.size = Pt(10.5)
    doc.add_paragraph()

    # ============== 3. GLOSARIO: QUE SIGNIFICA PAUSAR ==============
    add_heading(doc, "3. Glosario — Que significa exactamente cada accion", level=1)

    add_heading(doc, "3.1 'PAUSAR': 4 formas posibles en ML", level=2)
    add_para(doc, "En MercadoLibre hay 4 formas distintas de 'dejar de promover un item'. Cada una tiene efecto distinto. Aclarar cual usamos es critico:",
             size=10.5, italic=True)

    pausar_table = [
        ["1. PAUSAR PUBLICACION (LO QUE HAREMOS)",
         'PUT /items/{id} {"status":"paused"}',
         "Item desaparece del buscador ML. Nadie lo ve. No genera ventas. No paga comision (no hay venta).",
         "100% reversible. Mantiene historial, ventas pasadas, reputacion, ranking SEO interno. Es lo MAS conservador."],
        ["2. SALIR DEL CATALOGO (ALTERNATIVA)",
         'PUT /items/{id} {"catalog_listing":false}',
         "Item sigue VISIBLE en busqueda como producto propio. NO compite por buy box en la ficha compartida.",
         "Reversible. NO resuelve la canibalizacion interna — sigue compitiendo en busqueda."],
        ["3. CERRAR PUBLICACION",
         'PUT /items/{id} {"status":"closed"}',
         "Item se cierra DEFINITIVAMENTE en ese item_id. Sin reactivacion posible.",
         "IRREVERSIBLE practico. Pierdes historial, reputacion del item. NO LO USAREMOS."],
        ["4. PAUSAR PUBLICIDAD (Product Ads)",
         "Endpoint ML Ads (otro)",
         "Solo aplica si el item tiene anuncios pagados. Pausa solo la publicidad, no el item.",
         "Reversible. NO APLICA: tus cuentas tienen ML Ads en VIEWER, bloqueado por API."],
    ]
    t = add_table_simple(doc,
                        ["Forma", "API call", "Efecto visible", "Reversibilidad / Notas"],
                        pausar_table, widths=[4.5, 4.5, 4, 4])
    # marcar la fila #1 en verde (la elegida)
    shade_cell(t.rows[1].cells[0], "C5E1A5")
    doc.add_paragraph()

    add_heading(doc, "Por que elegimos 'pausar publicacion' (forma 1)", level=3)
    razones = [
        "Los 28 casos son items donde otra cuenta tuya ya gana el buy box. Pausar el perdedor NO quita ventas (ya iban al ganador interno).",
        "Es 100% reversible: si maniana cambian las condiciones, reactivar es 1 API call.",
        "'Salir del catalogo' (#2) no resuelve canibalizacion interna — el item perdedor seguiria visible y consumiendo atencion del algoritmo.",
        "'Cerrar' (#3) es irreversible y pierdes historial del item. Mucho riesgo para una decision exploratoria.",
        "'Pausar publicidad' (#4) solo aplica si tuvieras ML Ads en esos items (no aplica — ML Ads en VIEWER).",
    ]
    for r in razones:
        doc.add_paragraph(r, style="List Bullet")
    doc.add_paragraph()

    add_heading(doc, "3.2 'BAJAR_PRECIO': que hace exactamente", level=2)
    bajar_lines = [
        'API call: PUT /items/{id} {"price": <nuevo_valor>}',
        "El cambio aplica INMEDIATAMENTE al item ML (no afecta Web, Tienda, Falabella, Paris, Walmart).",
        "Defontana mantiene su precio sugerido sin modificacion.",
        "IMPORTANTE: Si tu Defontana sincroniza precios hacia ML (push automatico), el cambio podria revertirse en la siguiente sincronizacion. Verificar antes de aplicar.",
        "100% reversible: PUT con el precio anterior lo restaura.",
        "ML actualiza el buy box dentro de 1-6 horas tipicamente.",
    ]
    for line in bajar_lines:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph()

    # ============== 4. PLAN DETALLADO ==============
    doc.add_page_break()
    add_heading(doc, "4. Plan de accion (32 acciones pre-aprobadas)", level=1)

    # tabla resumen por cuenta
    add_heading(doc, "4.1 Distribucion por cuenta", level=2)
    plan_rows = [
        ["C1 (PREMIUMGRIFERIAS1)", total_por_cuenta.get("C1", 0),
         accion_cuenta.get(("PAUSAR", "C1"), 0), accion_cuenta.get(("BAJAR_PRECIO", "C1"), 0)],
        ["C2 (VICTTORINOOFICIAL2)", total_por_cuenta.get("C2", 0),
         accion_cuenta.get(("PAUSAR", "C2"), 0), accion_cuenta.get(("BAJAR_PRECIO", "C2"), 0)],
        ["C3 (NOVAGRIFERIAS3)", total_por_cuenta.get("C3", 0),
         accion_cuenta.get(("PAUSAR", "C3"), 0), accion_cuenta.get(("BAJAR_PRECIO", "C3"), 0)],
        ["TOTAL", len(acciones_si), 28, 4],
    ]
    t = add_table_simple(doc, ["Cuenta", "Total acciones", "PAUSAR", "BAJAR_PRECIO"],
                        plan_rows, widths=[6, 3, 3, 3])
    # ultima fila negrita
    for cell in t.rows[-1].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
        shade_cell(cell, "FFF59D")
    doc.add_paragraph()

    # los 4 BAJAR_PRECIO
    add_heading(doc, "4.2 Los 4 BAJAR_PRECIO (detalle)", level=2)
    bajar = sorted([a for a in acciones_si if a["accion"] == "BAJAR_PRECIO"],
                   key=lambda x: -(x.get("precio_actual") or 0))
    bajar_rows = []
    for a in bajar:
        bajar_rows.append([
            a["cuenta"],
            a["item_id"],
            (a.get("title") or "")[:40],
            fmt(a.get("precio_actual")),
            fmt(a.get("precio_sugerido")),
            f"{a.get('margen_actual_pct')}% -> {a.get('margen_proyectado_pct')}%",
        ])
    add_table_simple(doc,
                    ["Cuenta", "Item ID", "Producto", "Precio actual", "Precio nuevo", "Margen"],
                    bajar_rows, widths=[1.5, 2.5, 5.5, 2.5, 2.5, 3])
    doc.add_paragraph()

    # los 28 PAUSAR - resumen agrupado por cuenta
    add_heading(doc, "4.3 Los 28 PAUSAR (top 10 por valor)", level=2)
    pausar = sorted([a for a in acciones_si if a["accion"] == "PAUSAR"],
                   key=lambda x: -(x.get("precio_actual") or 0))[:10]
    pausar_rows = []
    for a in pausar:
        pausar_rows.append([
            a["cuenta"],
            a["item_id"],
            (a.get("title") or "")[:45],
            fmt(a.get("precio_actual")),
            f"{a.get('diff_pct_vs_ganador')}%",
            str(a.get("ganador_actual", ""))[:18],
        ])
    add_table_simple(doc,
                    ["Cuenta", "Item ID", "Producto", "Precio", "Mas caro", "Pierde vs"],
                    pausar_rows, widths=[1.5, 2.5, 5.5, 2.2, 1.8, 4])
    add_para(doc, "(Los 18 restantes estan en el Excel adjunto, hoja '3_ACCIONES_DETALLE')",
             size=9, italic=True, color="64748B")
    doc.add_paragraph()

    # ============== 5. PROYECCION ==============
    doc.add_page_break()
    add_heading(doc, "5. Proyeccion de impacto", level=1)

    add_para(doc,
             "Supuestos de venta usados (editables — los numeros reales pueden variar):",
             size=10.5)
    sup_rows = [
        ["ventas/mes si gana buy box", str(VENTAS_GANA), "Estimacion conservadora"],
        ["ventas/mes si pierde buy box", str(VENTAS_PIERDE), "Clicks/conversiones residuales"],
        ["Cobro canal ML", f"{int(ML_CHANNEL_FEE*100)}%", "Tu reporte de canales 2026-05-26"],
        ["IVA", f"{int(IVA*100)}%", "Estandar Chile"],
        ["Factor neto efectivo", f"{NETO_FACTOR:.4f}",
         "Lo que queda del precio publicado tras IVA + comision"],
    ]
    add_table_simple(doc, ["Parametro", "Valor", "Fuente"], sup_rows, widths=[6, 3, 6])
    doc.add_paragraph()

    add_heading(doc, "5.1 Impacto mensual y anual estimado", level=2)
    proy_rows = [
        ["Margen actual de items afectados", fmt(total_margen_actual), fmt(total_margen_actual * 12)],
        ["Margen proyectado en mismos items", fmt(total_margen_proy), fmt(total_margen_proy * 12)],
        ["Margen transferido a cuenta ganadora interna", fmt(total_margen_transf), fmt(total_margen_transf * 12)],
        ["IMPACTO NETO TOTAL", fmt(impacto_mes), fmt(impacto_mes * 12)],
    ]
    t = add_table_simple(doc, ["Concepto", "Mensual", "Anual"], proy_rows, widths=[8, 3.5, 3.5])
    # marcar ultima fila
    for cell in t.rows[-1].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
        shade_cell(cell, "FFF59D")
    doc.add_paragraph()

    add_para(doc, "ATENCION: Estos numeros son ORDEN DE MAGNITUD basados en supuestos.",
             bold=True, size=10.5, color="C62828")
    add_para(doc,
             "Si tienes datos reales de ventas/mes por SKU, te puedo recalcular con precision. "
             "El IMPACTO NETO es modesto en pesos absolutos porque solo afecta a 32 items "
             "(6% del catalogo). El valor real esta en simplificar el ecosistema y dejar de "
             "pagar comision Premium para perder contra ti mismo.",
             size=10.5, italic=True)
    doc.add_paragraph()

    # ============== 6. NO SE TOCA ==============
    add_heading(doc, "6. Lo que NO se toca (decisiones pendientes)", level=1)
    no_toca_rows = [
        ["13 items vendiendose a perdida",
         "URGENTE - revisar caso por caso. Posibles causas: error precio o costo desactualizado."],
        ["41 PIERDE_VS_EXT sin chance",
         "Decision: salir del catalogo, aceptar perdida, o migrar trafico a Web (72% margen vs 30% ML)."],
        ["26 BAJAR_PRECIO PENDIENTES (margen final <15%)",
         "Revisar 1 a 1 en Excel. Marca SI/NO en columna 'aprobado' del Excel."],
        ["5 sellers externos canibalizandote",
         "FACTORYNETCL, VIBRA TOOLS, OUTLETSEVENCL, MDWIS, GRIFERIALDAYCL. ¿Son distribuidores tuyos? Si SI: politica RRP."],
        ["112 fichas con stock 0 ('precio bloqueo $400.000')",
         "Operacional. ¿Reponen pronto? Mantener. ¿Discontinuados? Cerrar."],
        ["74 codigos Defontana sin publicacion ML",
         "Oportunidad: ¿son productos para listar? ¿O son B2B/internos?"],
        ["1 lookalike real (Kit Estanque WC)",
         "Reportar a soporte ML (brand protection) para que separe la ficha."],
    ]
    add_table_simple(doc, ["Pendiente", "Recomendacion"], no_toca_rows, widths=[6, 9])
    doc.add_paragraph()

    # ============== 7. COMO EJECUTAR ==============
    add_heading(doc, "7. Como ejecutar el plan", level=1)
    pasos = [
        "Revisar las 32 acciones aprobadas en el Excel `informe_revision_2026-05-26.xlsx` (hoja 3).",
        "Si quieres modificar alguna fila pre-aprobada: editar columna 'aprobado' (SI/NO/PENDIENTE).",
        "Ejecutar DRY-RUN primero: `python aplicar_plan_accion.py` (no toca ML).",
        "Verificar que dry-run muestra lo correcto.",
        "Ejecutar de verdad: `python aplicar_plan_accion.py --apply`",
        "El script crea backup en `backups/` y log CSV en `logs/`.",
        "Si algo sale mal: usar el backup para revertir (PUT status o PUT price con valores originales).",
    ]
    for i, paso in enumerate(pasos, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{i}. ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(paso)
        r2.font.size = Pt(11)

    doc.add_paragraph()
    add_para(doc, "Todas las acciones (PAUSAR y BAJAR_PRECIO) son 100% reversibles. El backup automatico permite revertir item por item si algo sale mal.",
             bold=True, italic=True, color="2E7D32", size=10.5)

    doc.save(DST)
    fecha = time.strftime("%Y-%m-%d")
    print(f"Word generado: {DST}")
    print(f"Excel:         {ROOT / f'informe_revision_{fecha}.xlsx'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
